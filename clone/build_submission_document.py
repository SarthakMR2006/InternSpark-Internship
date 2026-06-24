from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Spotify_Web_Player_UI_Clone_Submission.docx"
LIVE_SHOT = ROOT / "live-demo-screenshot.png"
REPO_SHOT = ROOT / "github-repository-screenshot.png"

LIVE_LINK = "https://sarthakmr2006.github.io/spotify-web-player-clone/"
SOURCE_LINK = "https://github.com/SarthakMR2006/spotify-web-player-clone"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.2)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F4F6")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(code.splitlines()):
        if i:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(31, 41, 55)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in [
        ("Heading 1", 17, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def build_doc():
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Spotify Web Player UI Clone")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Internship Project Submission Document")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "This project is a responsive UI-only clone of the Spotify Web Player. "
        "It was created using HTML, CSS, and JavaScript to demonstrate front-end "
        "layout, styling, responsiveness, and small interactive UI behavior."
    )

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.7), Inches(4.8)]
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = widths[index]
    rows = [
        ("Task", "Clone the UI of a real-world website using HTML, CSS, and JS"),
        ("Selected Option", "Spotify Web Player (UI only)"),
        ("Live Demo Link", LIVE_LINK),
        ("Source Code Link", SOURCE_LINK),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_shading(row.cells[0], "E8EEF5")
        row.cells[1].text = ""
        para = row.cells[1].paragraphs[0]
        if value.startswith("https://"):
            add_hyperlink(para, value, value)
        else:
            para.add_run(value)

    doc.add_heading("2. Hosted Links", level=1)
    p = doc.add_paragraph()
    p.add_run("Live Demo: ").bold = True
    add_hyperlink(p, LIVE_LINK, LIVE_LINK)
    p = doc.add_paragraph()
    p.add_run("GitHub Repository: ").bold = True
    add_hyperlink(p, SOURCE_LINK, SOURCE_LINK)

    doc.add_heading("3. Screenshots", level=1)
    doc.add_paragraph("Live demo screenshot:")
    if LIVE_SHOT.exists():
        doc.add_picture(str(LIVE_SHOT), width=Inches(6.4))
    doc.add_paragraph("GitHub repository screenshot:")
    if REPO_SHOT.exists():
        doc.add_picture(str(REPO_SHOT), width=Inches(6.4))

    doc.add_heading("4. Technology Used", level=1)
    for item in [
        "HTML for the page structure and content sections.",
        "CSS for Spotify-style layout, colors, spacing, responsive design, and player styling.",
        "JavaScript for UI-only interactions such as play button state, track selection, and mobile sidebar toggle.",
        "GitHub for source code hosting.",
        "GitHub Pages for live deployment.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("5. Features Implemented", level=1)
    for item in [
        "Spotify-style sidebar with navigation and playlist links.",
        "Top navigation bar with search input and account controls.",
        "Featured playlist hero section.",
        "Quick playlist cards and album recommendation cards.",
        "Bottom music player UI with play controls, progress bar, and volume control.",
        "Responsive layout for desktop and smaller screens.",
        "Interactive track selection that updates the now-playing area.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("6. Commands Used", level=1)
    doc.add_paragraph("The project can be tested locally with a simple static server command:")
    add_code_block(doc, "python -m http.server 4173 --bind 127.0.0.1")
    doc.add_paragraph("Then open the local preview URL:")
    add_code_block(doc, "http://127.0.0.1:4173/")
    doc.add_paragraph("GitHub Pages deployment was completed from the GitHub website:")
    for step in [
        "Open the repository on GitHub.",
        "Go to Settings > Pages.",
        "Choose Source: Deploy from a branch.",
        "Choose Branch: main and Folder: / (root).",
        "Save the settings and use the generated github.io live demo link.",
    ]:
        add_numbered(doc, step)

    doc.add_heading("7. Code Snippets", level=1)
    doc.add_heading("HTML structure", level=2)
    add_code_block(
        doc,
        """<link rel="stylesheet" href="styles.css" />
<script src="script.js" defer></script>

<aside class="sidebar" id="sidebar">
  <div class="brand">
    <span class="brand-mark"></span>
    <span>Spotify</span>
  </div>
</aside>

<footer class="player">
  <strong id="trackTitle">Today's Top Hits</strong>
</footer>""",
    )

    doc.add_heading("CSS layout and responsive design", level=2)
    add_code_block(
        doc,
        """.app-shell {
  display: grid;
  grid-template-columns: 278px minmax(0, 1fr);
  grid-template-rows: minmax(0, 1fr) 96px;
  gap: 8px;
  height: 100vh;
}

@media (max-width: 980px) {
  .sidebar {
    position: fixed;
    transform: translateX(-110%);
  }
  .sidebar.open {
    transform: translateX(0);
  }
}""",
    )

    doc.add_heading("JavaScript interaction", level=2)
    add_code_block(
        doc,
        """function setTrack(title) {
  trackTitle.textContent = title;
  isPlaying = true;
  mainPlay.textContent = "Pause";
  document.title = `${title} - Spotify Web Player UI Clone`;
}

menuToggle.addEventListener("click", () => {
  sidebar.classList.toggle("open");
});""",
    )

    doc.add_heading("8. UI-only Scope", level=1)
    doc.add_paragraph(
        "This submission is intentionally UI-only. It recreates the visual layout and front-end "
        "interactions of a Spotify-style web player, but it does not include real Spotify login, "
        "music streaming, backend APIs, real user accounts, or real playlist data."
    )

    doc.add_heading("9. Conclusion", level=1)
    doc.add_paragraph(
        "The project satisfies the requirement to clone the UI of a real-world website using "
        "HTML, CSS, and JavaScript. The live demo is hosted publicly through GitHub Pages, and "
        "the source code is available in a public GitHub repository."
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
